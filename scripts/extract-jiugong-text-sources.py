from __future__ import annotations

import json
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

SOURCE_ROOT = Path("/Users/guangmingxishe/LifeOS/20-29 工作事业/24.01 灵魂解码运营/九宫人生解码")
REPO_ROOT = Path(__file__).resolve().parents[1]
INVENTORY = REPO_ROOT / "docs/jiugong-sources/inventory.json"
EXTRACTED_ROOT = Path("/private/tmp/jiugong-source-extracted")
SUPPORTED = {"doc", "docx", "xls", "pdf", "md"}


def normalize(text: str) -> str:
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def extract_docx(source: Path) -> str:
    document = Document(source)
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.append("\t".join(cell.text for cell in row.cells))
    return normalize("\n".join(chunks))


def extract_doc(source: Path) -> str:
    result = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(source)],
        check=True,
        capture_output=True,
    )
    return normalize(result.stdout.decode("utf-8", errors="replace"))


def extract_pdf(source: Path) -> str:
    return normalize("\n".join(page.extract_text() or "" for page in PdfReader(source).pages))


def extract_xls(source: Path) -> str:
    with tempfile.TemporaryDirectory(prefix="jiugong-xls-") as temp_dir:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "xlsx", "--outdir", temp_dir, str(source)],
            check=True,
            capture_output=True,
        )
        converted = Path(temp_dir) / f"{source.stem}.xlsx"
        workbook = load_workbook(converted, read_only=True, data_only=False)
        chunks: list[str] = []
        for sheet in workbook.worksheets:
            chunks.append(f"## SHEET: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(values):
                    chunks.append("\t".join(values))
        return normalize("\n".join(chunks))


def extract(source: Path, file_format: str) -> tuple[str, str]:
    if file_format == "docx":
        return extract_docx(source), "python-docx"
    if file_format == "doc":
        return extract_doc(source), "textutil"
    if file_format == "pdf":
        return extract_pdf(source), "pypdf"
    if file_format == "xls":
        return extract_xls(source), "libreoffice+openpyxl"
    return normalize(source.read_text(encoding="utf-8")), "utf-8"


def main() -> None:
    payload = json.loads(INVENTORY.read_text(encoding="utf-8"))
    EXTRACTED_ROOT.mkdir(parents=True, exist_ok=True)
    failures: list[dict[str, str]] = []
    for item in payload["sources"]:
        if item["format"] not in SUPPORTED:
            continue
        source = SOURCE_ROOT / item["path"]
        try:
            text, extractor = extract(source, item["format"])
            (EXTRACTED_ROOT / f"{item['sourceId']}.txt").write_text(text, encoding="utf-8")
            item["status"] = "extracted" if text else "empty-needs-review"
            item["extractor"] = extractor
            item["characters"] = len(text)
            item["summary"] = re.sub(r"\s+", " ", text[:240])
        except Exception as error:  # status must be explicit; never pretend success
            item["status"] = "failed-needs-review"
            item["extractor"] = ""
            item["characters"] = 0
            item["summary"] = ""
            failures.append({"sourceId": item["sourceId"], "error": str(error)})
    INVENTORY.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({"processed": sum(1 for i in payload["sources"] if i["format"] in SUPPORTED), "failures": failures}, ensure_ascii=False))


if __name__ == "__main__":
    main()
